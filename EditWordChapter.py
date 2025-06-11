import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
import re
import win32com.client
import os
import logging
import sys

# 配置日志
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("word_editor.log"),
        logging.StreamHandler(sys.stdout)
    ]
)

class WordEditorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("章节删除程序")
        self.root.geometry("1200x800")

        # 首先创建日志窗口
        self.log_text = tk.Text(root, height=10)
        self.log_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        self.log_text.insert(tk.END, "日志输出区域...\n")
        self.log_text.configure(state='disabled')
        
        # 然后创建带滚动条的Treeview
        tree_frame = ttk.Frame(root)
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        tree_scroll = ttk.Scrollbar(tree_frame)
        tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.tree = ttk.Treeview(tree_frame, yscrollcommand=tree_scroll.set, selectmode='browse')
        self.tree.pack(fill=tk.BOTH, expand=True)
        tree_scroll.config(command=self.tree.yview)
        
        # 添加列
        self.tree["columns"] = ("index", "level", "position")
        self.tree.column("#0", width=400)
        self.tree.column("index", width=50)
        self.tree.column("level", width=50)
        self.tree.column("position", width=50)
        self.tree.heading("#0", text="章节标题")
        self.tree.heading("index", text="索引")
        self.tree.heading("level", text="级别")
        self.tree.heading("position", text="位置")

        # 按钮框架
        btn_frame = ttk.Frame(root)
        btn_frame.pack(fill=tk.X, padx=10, pady=5)
        
        self.btn = tk.Button(btn_frame, text="删除选中章节", command=self.delete_section)
        self.btn.pack(side=tk.LEFT, padx=5)
        
        self.refresh_btn = tk.Button(btn_frame, text="刷新文档", command=self.refresh_document)
        self.refresh_btn.pack(side=tk.LEFT, padx=5)
        
        self.debug_btn = tk.Button(btn_frame, text="调试信息", command=self.show_debug_info)
        self.debug_btn.pack(side=tk.LEFT, padx=5)

        # 状态栏
        self.status = tk.Label(root, text="就绪", bd=1, relief=tk.SUNKEN, anchor=tk.W)
        self.status.pack(side=tk.BOTTOM, fill=tk.X)

        # 初始化文档
        self.doc_path = os.path.abspath("word.docx")
        self.refresh_document()

    def log(self, message):
        """向日志窗口和文件添加日志"""
        self.log_text.configure(state='normal')
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
        self.log_text.configure(state='disabled')
        logging.info(message)

    def show_debug_info(self):
        """显示调试信息"""
        debug_info = "===== 当前文档调试信息 =====\n"
        debug_info += f"文档路径: {self.doc_path}\n"
        debug_info += f"段落总数: {len(self.doc.paragraphs)}\n"
        debug_info += f"标题数量: {len(self.headings)}\n\n"
        
        debug_info += "===== 前10个段落内容 =====\n"
        for i, para in enumerate(self.doc.paragraphs[:10]):
            debug_info += f"{i}: {para.text[:50]}{'...' if len(para.text) > 50 else ''}\n"
        
        debug_info += "\n===== 标题结构 =====\n"
        for i, (level, text, para, doc_index, number) in enumerate(self.headings[:20]):
            debug_info += f"{i}: L{level} - {number} {text[:30]}{'...' if len(text) > 30 else ''} (位置: {doc_index})\n"
        
        messagebox.showinfo("调试信息", debug_info)

    def refresh_document(self):
        """刷新文档内容"""
        self.status.config(text="正在加载文档...")
        self.root.update()
        self.log("开始加载文档...")
        
        try:
            self.doc = Document(self.doc_path)
            self.load_headings()
            self.status.config(text=f"已加载文档: {os.path.basename(self.doc_path)}")
            self.log(f"文档加载成功，共 {len(self.doc.paragraphs)} 个段落")
        except Exception as e:
            error_msg = f"加载文档失败: {str(e)}"
            messagebox.showerror("错误", error_msg)
            self.status.config(text=error_msg)
            self.log(error_msg)

    def load_headings(self):
        """加载文档标题结构"""
        self.log("开始加载标题结构...")
        self.headings = []
        self.tree.delete(*self.tree.get_children())
        
        # 使用COM接口获取带编号的标题
        self.log("调用COM接口获取带编号的标题...")
        numbered_paras = get_paragraph_numbering_from_word(self.doc_path)
        self.log(f"从COM接口获取到 {len(numbered_paras)} 个带编号的段落")
        
        # 创建段落索引映射
        para_index_map = {}
        for i, para in enumerate(self.doc.paragraphs):
            # 使用前50个字符作为键，避免超长文本问题
            key = para.text.strip()[:50]
            para_index_map[key] = (para, i)
        
        stack = []
        self.log("开始构建标题树结构...")
        
        # 创建根节点
        root_id = self.tree.insert("", "end", text="文档标题", values=(-1, 0, -1))
        stack.append((0, root_id))  # 0级节点作为根
        
        for i, (number, style, text) in enumerate(numbered_paras):
            if not number or not text:
                continue

            # 获取段落对象和索引
            key = text.strip()[:50]
            para_obj, doc_index = para_index_map.get(key, (None, -1))
            
            if para_obj is None:
                # 尝试完整匹配
                self.log(f"未找到匹配的段落 (短匹配): {key}")
                for para_text, (p_obj, p_idx) in para_index_map.items():
                    if text.strip() == para_text:
                        para_obj, doc_index = p_obj, p_idx
                        break
                if para_obj is None:
                    self.log(f"无法匹配段落: {text[:30]}...")
                    continue
                
            # 计算标题级别
            level = number.count('.') + 1 if '.' in number else 1
            self.headings.append((level, text, para_obj, doc_index, number))
            self.log(f"添加标题: {number} {text[:30]}... (级别: {level}, 位置: {doc_index})")

            # 插入到树结构
            parent_id = ""
            if stack:
                # 找到最近的父节点
                for j in range(len(stack)-1, -1, -1):
                    if stack[j][0] < level:
                        parent_id = stack[j][1]
                        break
            
            node_id = self.tree.insert(
                parent_id, "end", 
                text=f"{number} {text}", 
                values=(len(self.headings)-1, level, doc_index)
            )
            
            # 更新堆栈
            while stack and stack[-1][0] >= level:
                stack.pop()
            stack.append((level, node_id))
        
        self.log(f"标题结构加载完成，共 {len(self.headings)} 个标题")

    def delete_section(self):
        """删除选中的章节及其子章节"""
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("提示", "请先选择一个章节")
            return

        index = int(self.tree.item(selected[0], "values")[0])
        self.log(f"用户选择删除索引为 {index} 的章节")
        
        if index >= len(self.headings):
            error_msg = f"无效的选择: {index} (最大索引: {len(self.headings)-1})"
            messagebox.showerror("错误", error_msg)
            self.log(error_msg)
            return
            
        base_level, base_text, base_para, base_index, base_number = self.headings[index]
        self.log(f"选中的章节: {base_number} {base_text} (级别: {base_level}, 位置: {base_index})")
        
        # 找到要删除的起始位置
        start_index = base_index
        self.log(f"删除起始位置: {start_index}")
        
        # 找到结束位置 - 只删除当前章节及其子章节
        # 查找下一个同级或更高级别标题（不包括当前章节的子章节）
        end_index = len(self.doc.paragraphs)
        next_heading_found = False
        
        self.log("开始查找结束位置...")
        self.log(f"当前文档段落数: {len(self.doc.paragraphs)}")
        
        # 第一步：在当前章节之后查找下一个同级或更高级别标题
        for i in range(start_index + 1, len(self.doc.paragraphs)):
            # 检查位置i是否是标题
            is_heading = False
            heading_level = None
            
            # 通过位置索引匹配标题
            for h in self.headings:
                # 使用位置索引比较而非对象比较
                if h[3] == i:  # h[3]是标题在文档中的位置索引
                    is_heading = True
                    heading_level = h[0]
                    heading_text = h[1]
                    self.log(f"在位置 {i} 找到标题: {heading_text[:30]}... (级别: {heading_level})")
                    
                    if heading_level <= base_level:  # 同级或更高级别
                        end_index = i
                        next_heading_found = True
                        self.log(f"找到结束位置: {end_index} (标题级别: {heading_level} <= 基础级别: {base_level})")
                        break
            
            if next_heading_found:
                break
            elif not is_heading:
                # 记录非标题段落内容
                if i < start_index + 10 or i > len(self.doc.paragraphs) - 10:
                    self.log(f"位置 {i}: 非标题段落 - {self.doc.paragraphs[i].text[:50]}{'...' if len(self.doc.paragraphs[i].text) > 50 else ''}")
        
        # 如果没有找到后续标题，结束位置设为文档末尾
        if not next_heading_found:
            self.log("未找到后续标题，结束位置设为文档末尾")
            end_index = len(self.doc.paragraphs)

        self.log(f"最终删除范围: {start_index} 到 {end_index} (共 {end_index - start_index} 个段落)")

        # 删除范围内的段落（从后往前删除）
        self.status.config(text=f"正在删除章节: {base_text}...")
        self.root.update()
        self.log(f"开始删除 {start_index} 到 {end_index} 之间的段落...")
        
        # 记录要删除的段落内容
        for i in range(start_index, min(end_index, start_index + 5)):
            if i < len(self.doc.paragraphs):
                self.log(f"将删除段落 {i}: {self.doc.paragraphs[i].text[:50]}{'...' if len(self.doc.paragraphs[i].text) > 50 else ''}")
        
        if end_index - start_index > 5:
            self.log(f"... 省略 {end_index - start_index - 5} 个段落 ...")
        
        # 注意：只删除从start_index到end_index之间的内容
        # 不包括end_index位置的段落（即下一个章节的开始）
        delete_count = 0
        for i in range(end_index - 1, start_index - 1, -1):
            if i < len(self.doc.paragraphs):
                self._remove_paragraph(self.doc.paragraphs[i])
                delete_count += 1
        
        self.log(f"已删除 {delete_count} 个段落")
        
        # 保存文档
        try:
            self.log("正在保存文档...")
            self.doc.save(self.doc_path)
            success_msg = f"已删除章节: {base_text}"
            messagebox.showinfo("成功", success_msg)
            self.log(success_msg)
            self.refresh_document()
        except Exception as e:
            error_msg = f"保存失败: {str(e)}"
            messagebox.showerror("保存失败", error_msg)
            self.status.config(text=error_msg)
            self.log(error_msg)

    def _remove_paragraph(self, paragraph):
        """安全删除段落"""
        p = paragraph._element
        p.getparent().remove(p)


def get_paragraph_numbering_from_word(path):
    """使用COM接口获取带编号的段落 - 增强版"""
    logging.info("启动Word COM接口...")
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        logging.info(f"打开文档: {path}")
        doc = word.Documents.Open(path)
        
        result = []
        logging.info("开始遍历段落...")
        
        for para_idx, para in enumerate(doc.Paragraphs):
            try:
                text = para.Range.Text.strip()
                if not text:
                    continue
                    
                style = para.Style.NameLocal.strip()

                # 获取编号
                try:
                    number = para.Range.ListFormat.ListString.strip()
                except:
                    number = ""
                
                # 修复1：允许一级标题（带点）
                if number and number.endswith('.'):
                    # 去掉末尾的点，保留数字
                    number = number[:-1].strip()
                
                # 修复2：扩展标题识别规则
                is_valid_heading = False
                
                # 规则1：标准多级编号 (1, 1.1, 1.1.1)
                if number and re.fullmatch(r"\d+(\.\d+)*", number):
                    is_valid_heading = True
                
                # 规则2：一级标题 (1., 2., 3.)
                elif re.match(r"^\d+\.\s", text):
                    # 从文本中提取编号
                    match = re.match(r"^(\d+)\.", text)
                    if match:
                        number = match.group(1)
                        is_valid_heading = True
                
                # 规则3：带括号的编号 (1), 2), 3))
                elif re.match(r"^\d+\)\s", text):
                    match = re.match(r"^(\d+)\)", text)
                    if match:
                        number = match.group(1)
                        is_valid_heading = True
                
                # 规则4：中文数字编号（一、二、三）
                elif re.match(r"^[一二三四五六七八九十]+、", text):
                    match = re.match(r"^([一二三四五六七八九十]+)、", text)
                    if match:
                        number = match.group(1)
                        is_valid_heading = True
                
                if not is_valid_heading:
                    # 跳过图注
                    if re.match(r"^图\s*\d+", text):
                        continue
                    # 跳过短的非标题段落
                    if len(text) < 50 and not text.endswith(('：', ':')):
                        continue
                    
                    # 记录被跳过的段落
                    logging.debug(f"跳过段落: {text[:50]}{'...' if len(text) > 50 else ''}")
                    continue
                
                # 添加有效的标题
                result.append((number, style, text))
                logging.debug(f"添加标题: {number} | {text[:50]}{'...' if len(text) > 50 else ''}")
                
            except Exception as e:
                logging.error(f"处理段落 {para_idx} 时出错: {str(e)}")
        
        logging.info(f"共找到 {len(result)} 个带编号的段落")
        
        # 记录前10个找到的标题
        for i, (num, style, txt) in enumerate(result[:10]):
            logging.debug(f"标题{i+1}: {num} | {style} | {txt[:50]}{'...' if len(txt) > 50 else ''}")
        
        doc.Close(False)
        word.Quit()
        return result
    except Exception as e:
        logging.error(f"COM接口错误: {str(e)}")
        return []
    finally:
        # 确保释放COM对象
        try:
            del doc
            del word
        except:
            pass


if __name__ == "__main__":
    root = tk.Tk()
    app = WordEditorApp(root)
    root.mainloop()